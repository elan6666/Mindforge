"""Generated document artifact storage and export."""

from __future__ import annotations

from datetime import UTC, datetime
from functools import lru_cache
from io import BytesIO
import json
import os
from pathlib import Path
import re
from uuid import uuid4

from app.backend.core.config import Settings, get_settings
from app.backend.schemas.artifacts import ArtifactExportRequest, ArtifactSummary


class ArtifactService:
    """Generate and persist downloadable MD/PDF/DOCX/TEX artifacts."""

    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self.root = Path(settings.artifact_storage_path)
        self.root.mkdir(parents=True, exist_ok=True)
        self.index_path = self.root / "index.json"

    def export(self, payload: ArtifactExportRequest) -> ArtifactSummary:
        """Generate one artifact and persist it."""
        if len(payload.content) > self.settings.artifact_max_content_chars:
            raise ValueError(
                "Artifact source content is too large. "
                f"Limit is {self.settings.artifact_max_content_chars} characters."
            )
        artifact_id = str(uuid4())
        filename = f"{self._safe_filename(payload.title)}.{payload.format}"
        data = self._render(payload)
        artifact_path = self.root / f"{artifact_id}-{filename}"
        if not self._is_within(artifact_path, self.root):
            raise ValueError("Artifact path escapes the configured storage root.")
        artifact_path.write_bytes(data)
        summary = ArtifactSummary(
            artifact_id=artifact_id,
            title=payload.title,
            format=payload.format,
            filename=filename,
            mime_type=self._mime_type(payload.format),
            size_bytes=len(data),
            created_at=datetime.now(UTC).isoformat(),
            source_task_id=payload.source_task_id,
            download_url=f"/api/artifacts/{artifact_id}/download",
        )
        records = self._load()
        records[artifact_id] = {
            **summary.model_dump(mode="json"),
            "path": str(artifact_path),
        }
        self._save(records)
        return summary

    def list_artifacts(self) -> list[ArtifactSummary]:
        """Return generated artifact summaries."""
        return [
            ArtifactSummary.model_validate(
                {key: value for key, value in item.items() if key != "path"}
            )
            for item in self._load().values()
        ]

    def get_download(self, artifact_id: str) -> tuple[ArtifactSummary, Path] | None:
        """Return summary and path for a downloadable artifact."""
        record = self._load().get(artifact_id)
        if record is None:
            return None
        path = Path(str(record.get("path") or ""))
        if not self._is_within(path, self.root) or not path.exists():
            return None
        summary = ArtifactSummary.model_validate(
            {key: value for key, value in record.items() if key != "path"}
        )
        return summary, path

    def delete(self, artifact_id: str) -> bool:
        """Delete one generated artifact."""
        records = self._load()
        record = records.pop(artifact_id, None)
        if record is None:
            return False
        path = Path(str(record.get("path") or ""))
        if self._is_within(path, self.root) and path.exists():
            path.unlink()
        self._save(records)
        return True

    def _render(self, payload: ArtifactExportRequest) -> bytes:
        if payload.format == "md":
            return payload.content.encode("utf-8")
        if payload.format == "tex":
            return self._render_tex(payload).encode("utf-8")
        if payload.format == "docx":
            return self._render_docx(payload)
        if payload.format == "pdf":
            return self._render_pdf(payload)
        raise ValueError(f"Unsupported artifact format '{payload.format}'.")

    @staticmethod
    def _render_tex(payload: ArtifactExportRequest) -> str:
        escaped_title = payload.title.replace("&", r"\&").replace("%", r"\%")
        content = payload.content.replace("&", r"\&").replace("%", r"\%")
        return "\n".join(
            [
                r"\documentclass{article}",
                r"\usepackage[UTF8]{ctex}",
                r"\usepackage{geometry}",
                r"\geometry{margin=1in}",
                rf"\title{{{escaped_title}}}",
                r"\begin{document}",
                r"\maketitle",
                content,
                r"\end{document}",
                "",
            ]
        )

    @staticmethod
    def _render_docx(payload: ArtifactExportRequest) -> bytes:
        from docx import Document
        from docx.shared import Pt

        document = Document()
        document.core_properties.title = payload.title
        document.add_heading(payload.title, level=1)
        for line in payload.content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                document.add_heading(stripped[2:].strip(), level=1)
            elif stripped.startswith("## "):
                document.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith("### "):
                document.add_heading(stripped[4:].strip(), level=3)
            elif stripped.startswith("- "):
                document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            elif re.match(r"^\d+\.\s+", stripped):
                document.add_paragraph(
                    re.sub(r"^\d+\.\s+", "", stripped),
                    style="List Number",
                )
            else:
                document.add_paragraph(stripped)
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10.5)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _render_pdf(payload: ArtifactExportRequest) -> bytes:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = BytesIO()
        document = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=42,
            leftMargin=42,
            topMargin=48,
            bottomMargin=48,
            title=payload.title,
        )
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            base_font = "STSong-Light"
        except Exception:
            base_font = "Helvetica"
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "MindforgeTitle",
            parent=styles["Title"],
            fontName=base_font,
            fontSize=22,
            leading=30,
            textColor=colors.HexColor("#173f36"),
            spaceAfter=18,
        )
        heading_style = ParagraphStyle(
            "MindforgeHeading",
            parent=styles["Heading1"],
            fontName=base_font,
            fontSize=15,
            leading=21,
            textColor=colors.HexColor("#1f5f52"),
            spaceBefore=14,
            spaceAfter=8,
        )
        subheading_style = ParagraphStyle(
            "MindforgeSubheading",
            parent=styles["Heading2"],
            fontName=base_font,
            fontSize=12.5,
            leading=18,
            textColor=colors.HexColor("#5f4a27"),
            spaceBefore=10,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "MindforgeBody",
            parent=styles["BodyText"],
            fontName=base_font,
            fontSize=10,
            leading=16,
            spaceAfter=6,
        )
        bullet_style = ParagraphStyle(
            "MindforgeBullet",
            parent=body_style,
            leftIndent=16,
            firstLineIndent=-8,
        )
        story = [Paragraph(ArtifactService._escape_pdf_text(payload.title), title_style)]
        story.append(Spacer(1, 10))
        for line in payload.content.splitlines():
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 8))
                continue
            if stripped.startswith("# "):
                story.append(Paragraph(ArtifactService._escape_pdf_text(stripped[2:].strip()), heading_style))
            elif stripped.startswith("## "):
                story.append(Paragraph(ArtifactService._escape_pdf_text(stripped[3:].strip()), subheading_style))
            elif stripped.startswith("### "):
                story.append(Paragraph(ArtifactService._escape_pdf_text(stripped[4:].strip()), subheading_style))
            elif stripped.startswith("- "):
                story.append(Paragraph(f"• {ArtifactService._escape_pdf_text(stripped[2:].strip())}", bullet_style))
            elif re.match(r"^\d+\.\s+", stripped):
                story.append(Paragraph(ArtifactService._escape_pdf_text(stripped), bullet_style))
            else:
                story.append(Paragraph(ArtifactService._escape_pdf_text(stripped), body_style))
            story.append(Spacer(1, 6))
        document.build(story)
        return buffer.getvalue()

    @staticmethod
    def _escape_pdf_text(value: str) -> str:
        return (
            value.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    @staticmethod
    def _mime_type(format_name: str) -> str:
        return {
            "md": "text/markdown; charset=utf-8",
            "tex": "application/x-tex",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pdf": "application/pdf",
        }[format_name]

    @staticmethod
    def _safe_filename(value: str) -> str:
        safe = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "_", value).strip("._")
        return safe or "mindforge-artifact"

    def _load(self) -> dict[str, dict[str, object]]:
        if not self.index_path.exists():
            return {}
        return json.loads(self.index_path.read_text(encoding="utf-8")).get("artifacts", {})

    def _save(self, records: dict[str, dict[str, object]]) -> None:
        temp_path = self.index_path.with_suffix(self.index_path.suffix + ".tmp")
        temp_path.write_text(
            json.dumps({"artifacts": records}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        os.replace(temp_path, self.index_path)

    @staticmethod
    def _is_within(path: Path, root: Path) -> bool:
        try:
            path.resolve().relative_to(root.resolve())
            return True
        except ValueError:
            return False


@lru_cache(maxsize=1)
def get_artifact_service() -> ArtifactService:
    """Return cached artifact service."""
    return ArtifactService(get_settings())


def clear_artifact_service_cache() -> None:
    """Clear cached artifact service after settings changes."""
    get_artifact_service.cache_clear()
