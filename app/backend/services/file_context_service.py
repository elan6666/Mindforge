"""Uploaded file storage, parsing, chunking, and retrieval."""

from __future__ import annotations

from dataclasses import dataclass
from functools import lru_cache
from hashlib import sha256
from html import unescape
from io import BytesIO
import json
import os
from pathlib import Path
import re
from typing import Any
from uuid import uuid4

from app.backend.core.config import Settings, get_settings
from app.backend.schemas.file_context import (
    FileChunk,
    FileContextSummary,
    UploadedFileSummary,
)

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".csv",
    ".tsv",
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".css",
    ".html",
    ".htm",
    ".xml",
    ".yaml",
    ".yml",
    ".toml",
    ".log",
}
WORD_PATTERN = re.compile(r"[\w\u4e00-\u9fff]+", re.UNICODE)
MAX_STORED_CHUNKS = 240
CHUNK_SIZE = 1600
CHUNK_OVERLAP = 180
TEXT_EXCERPT_LIMIT = 2400


@dataclass(slots=True)
class StoredFileRecord:
    """Internal on-disk file record."""

    summary: UploadedFileSummary
    raw_path: str
    chunks: list[FileChunk]


class FileContextService:
    """Persist uploaded files and retrieve task-relevant chunks."""

    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self.root = Path(settings.file_storage_path)
        self.raw_dir = self.root / "raw"
        self.index_path = self.root / "index.json"
        self.root.mkdir(parents=True, exist_ok=True)
        self.raw_dir.mkdir(parents=True, exist_ok=True)

    def save_upload(
        self,
        *,
        filename: str,
        content_type: str | None,
        data: bytes,
    ) -> UploadedFileSummary:
        """Save, parse, chunk, and index one uploaded file."""
        file_id = str(uuid4())
        safe_name = self._safe_filename(filename or "upload.bin")
        file_dir = self.raw_dir / file_id
        file_dir.mkdir(parents=True, exist_ok=True)
        raw_path = file_dir / safe_name
        raw_path.write_bytes(data)

        parsed_text = ""
        parser = "unsupported"
        status = "unsupported"
        error_message: str | None = None
        try:
            parsed_text, parser = self._extract_text(
                data,
                filename=safe_name,
                content_type=content_type,
            )
            if parser == "unsupported":
                status = "unsupported"
            else:
                status = "parsed" if parsed_text.strip() else "empty"
        except Exception as exc:  # pragma: no cover - parser libraries vary by file
            status = "failed"
            error_message = str(exc)

        chunks = self._chunk_text(file_id, parsed_text)
        digest = sha256(data).hexdigest()
        summary = UploadedFileSummary(
            file_id=file_id,
            name=safe_name,
            mime_type=content_type,
            size_bytes=len(data),
            sha256=digest,
            status=status,
            parser=parser,
            text_excerpt=parsed_text[:TEXT_EXCERPT_LIMIT],
            char_count=len(parsed_text),
            chunk_count=len(chunks),
            error_message=error_message,
            metadata={},
        )
        self._upsert_record(
            StoredFileRecord(summary=summary, raw_path=str(raw_path), chunks=chunks)
        )
        return self._public_summary(summary)

    def list_files(self) -> list[UploadedFileSummary]:
        """Return all uploaded file summaries."""
        records = self._load_records()
        return [self._public_summary(record.summary) for record in records.values()]

    def get_file(self, file_id: str) -> UploadedFileSummary | None:
        """Return one uploaded file summary."""
        record = self._load_records().get(file_id)
        return self._public_summary(record.summary) if record else None

    def delete_file(self, file_id: str) -> bool:
        """Delete one uploaded file and its indexed chunks."""
        records = self._load_records()
        record = records.pop(file_id, None)
        if record is None:
            return False
        self._write_records(records)
        raw_path = Path(record.raw_path)
        if not self._is_within(raw_path, self.raw_dir):
            return True
        if raw_path.exists():
            raw_path.unlink()
        if raw_path.parent.exists() and raw_path.parent != self.raw_dir:
            try:
                raw_path.parent.rmdir()
            except OSError:
                pass
        return True

    def resolve_context(
        self,
        *,
        file_ids: list[str],
        query: str,
        limit: int = 8,
    ) -> FileContextSummary:
        """Retrieve relevant chunks for the prompt from uploaded files."""
        if not file_ids:
            return FileContextSummary(status="skipped")
        records = self._load_records()
        selected_files: list[UploadedFileSummary] = []
        candidate_chunks: list[FileChunk] = []
        warnings: list[str] = []
        seen: set[str] = set()
        for file_id in file_ids:
            if file_id in seen:
                continue
            seen.add(file_id)
            record = records.get(file_id)
            if record is None:
                warnings.append(f"Unknown uploaded file '{file_id}'.")
                continue
            selected_files.append(self._public_summary(record.summary))
            candidate_chunks.extend(record.chunks)

        ranked = self._rank_chunks(candidate_chunks, query)[:limit]
        status = "retrieved" if ranked else ("no_chunks" if selected_files else "not_found")
        return FileContextSummary(
            status=status,
            file_ids=list(seen),
            files=selected_files,
            chunks=ranked,
            warnings=warnings,
        )

    @classmethod
    def _extract_text(
        cls,
        data: bytes,
        *,
        filename: str,
        content_type: str | None,
    ) -> tuple[str, str]:
        suffix = Path(filename).suffix.lower()
        normalized_type = (content_type or "").split(";", 1)[0].lower()
        if normalized_type == "application/pdf" or suffix == ".pdf":
            return cls._extract_pdf(data), "pypdf"
        if (
            normalized_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or suffix == ".docx"
        ):
            return cls._extract_docx(data), "python-docx"
        if normalized_type in {"text/html", "application/xhtml+xml"} or suffix in {
            ".html",
            ".htm",
        }:
            return cls._extract_html(data), "beautifulsoup-html"
        if normalized_type.startswith("text/") or suffix in TEXT_EXTENSIONS:
            return cls._decode_text(data), "plain-text"
        return "", "unsupported"

    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        pages: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {index}]\n{text}")
        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        from docx import Document

        document = Document(BytesIO(data))
        parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(part for part in parts if part.strip())

    @staticmethod
    def _extract_html(data: bytes) -> str:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(FileContextService._decode_text(data), "html.parser")
        for tag in soup(["script", "style", "noscript", "svg"]):
            tag.decompose()
        title = soup.title.get_text(" ", strip=True) if soup.title else ""
        body = soup.get_text("\n", strip=True)
        return "\n\n".join(part for part in [title, body] if part)

    @staticmethod
    def _decode_text(data: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
            try:
                return unescape(data.decode(encoding))
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _chunk_text(file_id: str, text: str) -> list[FileChunk]:
        normalized = re.sub(r"\n{3,}", "\n\n", text.strip())
        if not normalized:
            return []
        chunks: list[FileChunk] = []
        start = 0
        order = 0
        while start < len(normalized) and len(chunks) < MAX_STORED_CHUNKS:
            end = min(len(normalized), start + CHUNK_SIZE)
            if end < len(normalized):
                paragraph_break = normalized.rfind("\n\n", start, end)
                if paragraph_break > start + 400:
                    end = paragraph_break
            chunk_text = normalized[start:end].strip()
            if chunk_text:
                chunks.append(
                    FileChunk(
                        chunk_id=f"{file_id}:{order}",
                        file_id=file_id,
                        order=order,
                        text=chunk_text,
                        char_start=start,
                        char_end=end,
                    )
                )
                order += 1
            if end >= len(normalized):
                break
            start = max(end - CHUNK_OVERLAP, start + 1)
        return chunks

    @staticmethod
    def _rank_chunks(chunks: list[FileChunk], query: str) -> list[FileChunk]:
        terms = [term.lower() for term in WORD_PATTERN.findall(query)]
        if not terms:
            return chunks[:8]
        unique_terms = set(terms)
        ranked: list[FileChunk] = []
        for chunk in chunks:
            text_lower = chunk.text.lower()
            score = 0.0
            for term in unique_terms:
                if term in text_lower:
                    score += text_lower.count(term) * (2.0 if len(term) > 1 else 0.4)
            if query.strip() and query.strip().lower() in text_lower:
                score += 5.0
            if score <= 0:
                continue
            ranked.append(chunk.model_copy(update={"score": score}))
        if not ranked:
            return chunks[:8]
        return sorted(ranked, key=lambda chunk: chunk.score, reverse=True)

    def _load_records(self) -> dict[str, StoredFileRecord]:
        if not self.index_path.exists():
            return {}
        payload = json.loads(self.index_path.read_text(encoding="utf-8"))
        records: dict[str, StoredFileRecord] = {}
        for file_id, item in payload.get("files", {}).items():
            summary = UploadedFileSummary.model_validate(item["summary"])
            chunks = [FileChunk.model_validate(chunk) for chunk in item.get("chunks", [])]
            records[file_id] = StoredFileRecord(
                summary=summary,
                raw_path=str(item.get("raw_path") or summary.metadata.get("raw_path") or ""),
                chunks=chunks,
            )
        return records

    def _write_records(self, records: dict[str, StoredFileRecord]) -> None:
        payload = {
            "files": {
                file_id: {
                    "summary": record.summary.model_dump(mode="json"),
                    "raw_path": record.raw_path,
                    "chunks": [chunk.model_dump(mode="json") for chunk in record.chunks],
                }
                for file_id, record in records.items()
            }
        }
        temp_path = self.index_path.with_suffix(self.index_path.suffix + ".tmp")
        temp_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        os.replace(temp_path, self.index_path)

    def _upsert_record(self, record: StoredFileRecord) -> None:
        records = self._load_records()
        records[record.summary.file_id] = record
        self._write_records(records)

    @staticmethod
    def _safe_filename(filename: str) -> str:
        safe = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "_", filename).strip("._")
        return safe or "upload.bin"

    @staticmethod
    def _public_summary(summary: UploadedFileSummary) -> UploadedFileSummary:
        """Return public file metadata without local filesystem paths."""
        metadata = dict(summary.metadata)
        metadata.pop("raw_path", None)
        return summary.model_copy(update={"metadata": metadata})

    @staticmethod
    def _is_within(path: Path, root: Path) -> bool:
        try:
            path.resolve().relative_to(root.resolve())
            return True
        except ValueError:
            return False


@lru_cache(maxsize=1)
def get_file_context_service() -> FileContextService:
    """Return cached uploaded file context service."""
    return FileContextService(get_settings())


def clear_file_context_service_cache() -> None:
    """Clear cached file context service after settings changes in tests."""
    get_file_context_service.cache_clear()
