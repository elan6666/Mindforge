from io import BytesIO

from docx import Document

from app.backend.core.config import Settings
from app.backend.services.file_context_service import FileContextService


def make_file_service(tmp_path) -> FileContextService:
    return FileContextService(
        Settings(
            openhands_mode="mock",
            sqlite_db_path=str(tmp_path / "mindforge.db"),
            file_storage_path=str(tmp_path / "files"),
        )
    )


def test_file_context_service_parses_text_and_retrieves_relevant_chunks(tmp_path):
    service = make_file_service(tmp_path)

    summary = service.save_upload(
        filename="notes.md",
        content_type="text/markdown",
        data=b"Alpha planning notes.\n\nBeta retrieval details for Mindforge.",
    )
    context = service.resolve_context(
        file_ids=[summary.file_id],
        query="retrieval details",
        limit=4,
    )

    assert summary.status == "parsed"
    assert summary.parser == "plain-text"
    assert summary.chunk_count == 1
    assert "raw_path" not in summary.metadata
    assert "raw_path" not in context.files[0].metadata
    assert context.status == "retrieved"
    assert context.chunks[0].file_id == summary.file_id
    assert "Beta retrieval details" in context.chunks[0].text


def test_file_context_service_parses_docx_tables_and_paragraphs(tmp_path):
    service = make_file_service(tmp_path)
    document = Document()
    document.add_paragraph("Reviewer standard: concise contribution claims.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Clarity"
    buffer = BytesIO()
    document.save(buffer)

    summary = service.save_upload(
        filename="paper-review.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        data=buffer.getvalue(),
    )
    context = service.resolve_context(
        file_ids=[summary.file_id],
        query="contribution clarity",
        limit=4,
    )

    assert summary.status == "parsed"
    assert summary.parser == "python-docx"
    assert summary.chunk_count == 1
    assert "Reviewer standard" in summary.text_excerpt
    assert "Metric | Clarity" in context.chunks[0].text


def test_file_context_service_reports_unsupported_binary(tmp_path):
    service = make_file_service(tmp_path)

    summary = service.save_upload(
        filename="archive.bin",
        content_type="application/octet-stream",
        data=b"\x00\x01\x02",
    )
    context = service.resolve_context(
        file_ids=[summary.file_id],
        query="anything",
        limit=4,
    )

    assert summary.status == "unsupported"
    assert summary.parser == "unsupported"
    assert summary.chunk_count == 0
    assert context.status == "no_chunks"
